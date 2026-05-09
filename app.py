
import streamlit as st
import pandas as pd
import json, uuid, base64
from pathlib import Path
from datetime import date, datetime
from io import BytesIO
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

APP_DIR = Path(__file__).parent
DATA_DIR = APP_DIR / "data"
ASSETS_DIR = APP_DIR / "assets"
CUSTOMERS_FILE = DATA_DIR / "customers.json"
DOCUMENTS_FILE = DATA_DIR / "documents.json"
SETTINGS_FILE = DATA_DIR / "settings.json"
LOGO_PATH = ASSETS_DIR / "logo.png"
STAMP_PATH = ASSETS_DIR / "stamp.png"

DATA_DIR.mkdir(exist_ok=True)
ASSETS_DIR.mkdir(exist_ok=True)

COMPANY = {
    "name": "CASA ARTE PRIVEE FZ-LLC",
    "address": "AL Hulaila Industrial Zone-FZ, Ras Al Khaimah, United Arab Emirates",
    "license": "LICENCE NO. 5038159",
    "phone": "+393333920771 / +971585911456",
    "email": "casaartepriveefze@artedicasaae.co",
}

BANKS = {
    "EUR": [
        {"Bank Name":"MASHREQ BANK PSC UAE","Currency":"EUR","IBAN":"AE830330000019102080621","SWIFT/BIC":"BOMLAEADXXX","Bank Address":"Umniyati Street, Burj Khalifa Community, Dubai, UAE. Postal Address: P.O. Box 1250, Dubai, UAE."},
        {"Bank Name":"WIO BANK PJSC","Currency":"EUR","IBAN":"AE830860000009168276489","SWIFT/BIC":"WIOBAEADXXX","Bank Address":"Etihad Airways Centre, 5th Floor, Abu Dhabi, UAE."},
    ],
    "USD": [
        {"Bank Name":"MASHREQ BANK PSC UAE","Currency":"USD","IBAN":"AE130330000019102080620","SWIFT/BIC":"BOMLAEADXXX","Bank Address":"Umniyati Street, Burj Khalifa Community, Dubai, UAE. Postal Address: P.O. Box 1250, Dubai, UAE."},
        {"Bank Name":"WIO BANK PJSC","Currency":"USD","IBAN":"AE480860000009878943739","SWIFT/BIC":"WIOBAEADXXX","Bank Address":"Etihad Airways Centre, 5th Floor, Abu Dhabi, UAE."},
    ],
    "AED": [
        {"Bank Name":"WIO BANK PJSC","Currency":"AED","IBAN":"AE730860000009886170371","SWIFT/BIC":"WIOBAEADXXX","Bank Address":"Etihad Airways Centre, 5th Floor, Abu Dhabi, UAE."},
    ],
}

PRODUCT_COLS = ["Brand","Product Details","Size","Finish","Qty","Rate Per Piece"]
PACK_COLS = ["Box No","Part","Brand","Product Details","Length","Breadth","Height","CBM","GW","NW"]

def load_json(path, default):
    if not path.exists():
        path.write_text(json.dumps(default, indent=2))
        return default
    try:
        return json.loads(path.read_text())
    except Exception:
        return default

def save_json(path, data):
    path.write_text(json.dumps(data, indent=2, ensure_ascii=False))

def img_b64(path):
    return base64.b64encode(path.read_bytes()).decode() if path.exists() else ""

def money(v, currency):
    symbol = {"EUR":"€","USD":"$","AED":"AED"}.get(currency, currency)
    return f"{symbol} {float(v or 0):,.2f}"

def number_prefix(doc_type):
    return "CAP/PI" if doc_type == "Proforma Invoice" else "CAP/INV"

def next_number(doc_type, docs):
    prefix = number_prefix(doc_type)
    year = datetime.now().year
    nums = []
    for d in docs:
        n = str(d.get("number",""))
        if n.startswith(f"{prefix}/{year}/"):
            try:
                nums.append(int(n.split("/")[-1]))
            except Exception:
                pass
    return f"{prefix}/{year}/{(max(nums) if nums else 0)+1:04d}"

def calculate(products, discount_type, discount_value, shipping_enabled, shipping_cost):
    subtotal = 0.0
    for p in products:
        subtotal += float(p.get("Qty",0) or 0) * float(p.get("Rate Per Piece",0) or 0)
    discount = subtotal * float(discount_value or 0) / 100 if discount_type == "Percentage" else float(discount_value or 0)
    shipping = float(shipping_cost or 0) if shipping_enabled else 0.0
    return subtotal, discount, shipping, subtotal - discount + shipping


def is_real_packing_row(row):
    return bool(str(row.get("Brand", "")).strip()) or bool(str(row.get("Product Details", "")).strip())

def clean_packing_rows(rows):
    cleaned = []
    for row in rows or []:
        if not is_real_packing_row(row):
            continue
        cleaned.append(row)
    for idx, row in enumerate(cleaned, 1):
        row["Box No"] = idx
    return cleaned

def packing_summary(rows):
    real_rows = clean_packing_rows(rows)
    return {
        "Total Boxes": len(real_rows),
        "Total CBM": round(sum(float(x.get("CBM", 0) or 0) for x in real_rows), 3),
        "Total GW": round(sum(float(x.get("GW", 0) or 0) for x in real_rows), 2),
        "Total NW": round(sum(float(x.get("NW", 0) or 0) for x in real_rows), 2),
    }

def packing_from_products(products, existing=None):
    existing = existing or []
    out = []

    # Existing packing rows are preserved so split boxes/parts are not lost on edit.
    for i, p in enumerate(products):
        matching = [
            row for row in existing
            if row.get("Product Details", "") == p.get("Product Details", "")
            and row.get("Brand", "") == p.get("Brand", "")
        ]

        if matching:
            for row in matching:
                l = float(row.get("Length", 0) or 0)
                b = float(row.get("Breadth", 0) or 0)
                h = float(row.get("Height", 0) or 0)
                box_qty = float(row.get("Box Qty", 1) or 1)
                row["Box No"] = int(row.get("Box No", len(out) + 1) or len(out) + 1)
                row["Part"] = row.get("Part", "1/1")
                row["Brand"] = p.get("Brand", "")
                row["Product Details"] = p.get("Product Details", "")
                row["CBM"] = round(l * b * h * box_qty / 1000000, 3)
                row["GW"] = float(row.get("GW", 0) or 0)
                row["NW"] = float(row.get("NW", 0) or 0)
                out.append(row)
        else:
            out.append({
                "Box No": len(out) + 1,
                "Part": "1/1",
                "Brand": p.get("Brand", ""),
                "Product Details": p.get("Product Details", ""),
                "Length": 0.0,
                "Breadth": 0.0,
                "Height": 0.0,
                "CBM": 0.0,
                "GW": 0.0,
                "NW": 0.0,
            })

    return clean_packing_rows(out)


def pdf_header_footer(canvas, doc, title=""):
    w, h = A4
    canvas.saveState()

    # Logo
    if LOGO_PATH.exists():
        canvas.drawImage(
            str(LOGO_PATH),
            12 * mm,
            h - 24 * mm,
            width=35 * mm,
            height=16 * mm,
            preserveAspectRatio=True,
            mask="auto"
        )

    # Header text
    canvas.setFillColor(colors.HexColor("#071c2e"))
    canvas.setFont("Helvetica-Bold", 11)
    canvas.drawRightString(w - 12 * mm, h - 15 * mm, COMPANY["name"])

    canvas.setFillColor(colors.HexColor("#9b763c"))
    canvas.setFont("Helvetica-Bold", 9)
    canvas.drawRightString(w - 12 * mm, h - 20 * mm, title)

    canvas.setStrokeColor(colors.HexColor("#d0aa65"))
    canvas.line(12 * mm, h - 27 * mm, w - 12 * mm, h - 27 * mm)

    # Stamp bottom right
    if STAMP_PATH.exists():
        try:
            canvas.saveState()
            canvas.setFillAlpha(0.55)
            canvas.drawImage(
                str(STAMP_PATH),
                w - 48 * mm,
                15 * mm,
                width=32 * mm,
                height=32 * mm,
                preserveAspectRatio=True,
                mask="auto"
            )
            canvas.restoreState()
        except Exception:
            canvas.drawImage(
                str(STAMP_PATH),
                w - 48 * mm,
                15 * mm,
                width=32 * mm,
                height=32 * mm,
                preserveAspectRatio=True,
                mask="auto"
            )

    # Footer
    canvas.setFillColor(colors.grey)
    canvas.setFont("Helvetica", 7)
    canvas.drawCentredString(
        w / 2,
        8 * mm,
        f"{COMPANY['name']} | {COMPANY['email']} | Page {canvas.getPageNumber()}"
    )

    canvas.restoreState()




ONES = ["", "One", "Two", "Three", "Four", "Five", "Six", "Seven", "Eight", "Nine", "Ten", "Eleven", "Twelve", "Thirteen", "Fourteen", "Fifteen", "Sixteen", "Seventeen", "Eighteen", "Nineteen"]
TENS = ["", "", "Twenty", "Thirty", "Forty", "Fifty", "Sixty", "Seventy", "Eighty", "Ninety"]

def number_to_words(n):
    n = int(n)
    if n == 0:
        return "Zero"
    if n < 20:
        return ONES[n]
    if n < 100:
        return TENS[n // 10] + ((" " + ONES[n % 10]) if n % 10 else "")
    if n < 1000:
        return ONES[n // 100] + " Hundred" + ((" " + number_to_words(n % 100)) if n % 100 else "")
    if n < 1000000:
        return number_to_words(n // 1000) + " Thousand" + ((" " + number_to_words(n % 1000)) if n % 1000 else "")
    if n < 1000000000:
        return number_to_words(n // 1000000) + " Million" + ((" " + number_to_words(n % 1000000)) if n % 1000000 else "")
    return number_to_words(n // 1000000000) + " Billion" + ((" " + number_to_words(n % 1000000000)) if n % 1000000000 else "")

def amount_in_words(amount, currency):
    amount = round(float(amount or 0), 2)
    whole = int(amount)
    cents = int(round((amount - whole) * 100))
    currency_name = {"EUR": "Euros", "USD": "US Dollars", "AED": "UAE Dirhams"}.get(currency, currency)
    minor_name = {"EUR": "Cents", "USD": "Cents", "AED": "Fils"}.get(currency, "Cents")
    words = f"{number_to_words(whole)} {currency_name}"
    if cents:
        words += f" and {number_to_words(cents)} {minor_name}"
    return words + " Only"


def build_excel(docdata):
    buffer = BytesIO()
    products = docdata.get("products", [])
    packing = packing_from_products(products, docdata.get("packing", []))
    subtotal, disc, shipc, total = calculate(
        products,
        docdata.get("discount_type", "Percentage"),
        docdata.get("discount_value", 0),
        docdata.get("shipping_enabled", False),
        docdata.get("shipping_cost", 0),
    )

    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        info_rows = [
            ["Document Type", docdata.get("type", "")],
            ["Document No", docdata.get("number", "")],
            ["Date", docdata.get("date", "")],
            ["Currency", docdata.get("currency", "")],
            ["Customer", docdata.get("bill_to", {}).get("Company Name", "")],
            ["Grand Total", total],
        ]
        pd.DataFrame(info_rows, columns=["Field", "Value"]).to_excel(writer, sheet_name="Document Info", index=False)

        invoice_rows = []
        for i, p in enumerate(products, 1):
            qty = float(p.get("Qty", 0) or 0)
            rate = float(p.get("Rate Per Piece", 0) or 0)
            invoice_rows.append({
                "SL": i,
                "Brand": p.get("Brand", ""),
                "Product Details": p.get("Product Details", ""),
                "Size": p.get("Size", ""),
                "Finish": p.get("Finish", ""),
                "Qty": qty,
                "Rate Per Piece": rate,
                "Total Qty": qty,
                "Amount": qty * rate,
            })
        pd.DataFrame(invoice_rows).to_excel(writer, sheet_name="Invoice", index=False)

        total_rows = [
            {"Description": "Subtotal", "Amount": subtotal},
            {"Description": "Discount", "Amount": disc},
            {"Description": "Shipping", "Amount": shipc},
            {"Description": "Grand Total", "Amount": total},
        ]
        pd.DataFrame(total_rows).to_excel(writer, sheet_name="Totals", index=False)

        pd.DataFrame(BANKS.get(docdata.get("currency", "EUR"), [])).to_excel(writer, sheet_name="Bank Details", index=False)

        pd.DataFrame([{"Terms & Conditions": docdata.get("terms", "")}]).to_excel(writer, sheet_name="Terms", index=False)

        if docdata.get("type") == "Invoice":
            pd.DataFrame(packing).to_excel(writer, sheet_name="Packing List", index=False)
            summary = [{
                "Total Boxes": len(packing),
                "Total CBM": sum(float(x.get("CBM", 0) or 0) for x in packing),
                "Total GW": sum(float(x.get("GW", 0) or 0) for x in packing),
                "Total NW": sum(float(x.get("NW", 0) or 0) for x in packing),
            }]
            pd.DataFrame(summary).to_excel(writer, sheet_name="Packing Summary", index=False)

        for sheet in writer.sheets.values():
            for col in sheet.columns:
                max_len = 12
                col_letter = col[0].column_letter
                for cell in col:
                    try:
                        max_len = max(max_len, len(str(cell.value or "")) + 2)
                    except Exception:
                        pass
                sheet.column_dimensions[col_letter].width = min(max_len, 45)

    buffer.seek(0)
    return buffer.getvalue()



def build_word(docdata):
    buffer = BytesIO()
    products = docdata.get("products", [])
    packing = packing_from_products(products, docdata.get("packing", []))
    subtotal, disc, shipc, total = calculate(
        products,
        docdata.get("discount_type", "Percentage"),
        docdata.get("discount_value", 0),
        docdata.get("shipping_enabled", False),
        docdata.get("shipping_cost", 0),
    )

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    if LOGO_PATH.exists():
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(LOGO_PATH), width=Inches(1.8))

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(docdata.get("type", "").upper())
    r.bold = True
    r.font.size = Pt(16)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"No: {docdata.get('number','')} | Date: {docdata.get('date','')} | Currency: {docdata.get('currency','')}").bold = True

    document.add_paragraph(f"{COMPANY['name']}\n{COMPANY['address']}\n{COMPANY['license']}\n{COMPANY['phone']}\n{COMPANY['email']}")

    bill = docdata.get("bill_to", {})
    ship = docdata.get("ship_to", {})
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.cell(0,0).text = "BILL TO\n" + "\n".join([
        bill.get("Company Name",""), bill.get("Registration Number",""), bill.get("GST/VAT",""),
        bill.get("Contact Person",""), bill.get("Phone",""), bill.get("Email",""),
        bill.get("Address",""), bill.get("Country","")
    ])
    table.cell(0,1).text = "SHIP TO\n" + ("Same as Bill To" if docdata.get("ship_same") else "\n".join([
        ship.get("Company Name",""), ship.get("Contact Person",""), ship.get("Phone",""),
        ship.get("Email",""), ship.get("Address",""), ship.get("Country","")
    ]))

    document.add_paragraph("")
    prod_table = document.add_table(rows=1, cols=8)
    prod_table.style = "Table Grid"
    headers = ["SL", "Brand", "Product Details", "Size", "Finish", "Qty", "Rate/PC", "Amount"]
    for i, h in enumerate(headers):
        prod_table.rows[0].cells[i].text = h
    for idx, p in enumerate(products, 1):
        qty = float(p.get("Qty", 0) or 0)
        rate = float(p.get("Rate Per Piece", 0) or 0)
        row = prod_table.add_row().cells
        vals = [idx, p.get("Brand",""), p.get("Product Details",""), p.get("Size",""), p.get("Finish",""), qty, money(rate, docdata.get("currency","EUR")), money(qty*rate, docdata.get("currency","EUR"))]
        for i, v in enumerate(vals):
            row[i].text = str(v)

    document.add_paragraph("")
    totals_table = document.add_table(rows=4, cols=2)
    totals_table.style = "Table Grid"
    rows = [
        ("Subtotal", money(subtotal, docdata.get("currency","EUR"))),
        ("Discount", f"- {money(disc, docdata.get('currency','EUR'))}"),
        ("Shipping", money(shipc, docdata.get("currency","EUR"))),
        ("Grand Total", money(total, docdata.get("currency","EUR"))),
    ]
    for i, (label, val) in enumerate(rows):
        totals_table.cell(i,0).text = label
        totals_table.cell(i,1).text = val

    p = document.add_paragraph()
    r = p.add_run("Amount in Words: ")
    r.bold = True
    p.add_run(amount_in_words(total, docdata.get("currency","EUR")))

    document.add_paragraph("")
    document.add_heading("Bank Details", level=2)
    bank_table = document.add_table(rows=1, cols=5)
    bank_table.style = "Table Grid"
    for i, h in enumerate(["Bank Name", "Currency", "IBAN", "SWIFT/BIC", "Bank Address"]):
        bank_table.rows[0].cells[i].text = h
    for b in BANKS.get(docdata.get("currency","EUR"), []):
        row = bank_table.add_row().cells
        for i, h in enumerate(["Bank Name", "Currency", "IBAN", "SWIFT/BIC", "Bank Address"]):
            row[i].text = b.get(h, "")

    document.add_page_break()
    if LOGO_PATH.exists():
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO_PATH), width=Inches(1.5))
    document.add_heading("Terms & Conditions", level=1)
    document.add_paragraph(docdata.get("terms", ""))

    document.add_paragraph("")
    sig = document.add_table(rows=2, cols=2)
    sig.style = "Table Grid"
    sig.cell(0,0).text = "Seller Signature"
    sig.cell(0,1).text = "Buyer Signature"
    sig.cell(1,0).text = "HARSH TEJPAL RANA\nOWNER\nCASA ARTE PRIVEE FZ-LLC"
    sig.cell(1,1).text = "Name:\nCompany:\nDate:"
    if STAMP_PATH.exists():
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run().add_picture(str(STAMP_PATH), width=Inches(1.0))

    if docdata.get("type") == "Invoice":
        document.add_page_break()
        if LOGO_PATH.exists():
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(LOGO_PATH), width=Inches(1.5))
        document.add_heading("Packing List", level=1)
        pack_table = document.add_table(rows=1, cols=10)
        pack_table.style = "Table Grid"
        for i, h in enumerate(["SL", "Box No", "Part", "Brand", "Product Details", "L", "B", "H", "CBM", "GW/NW"]):
            pack_table.rows[0].cells[i].text = h
        for i, p in enumerate(packing, 1):
            row = pack_table.add_row().cells
            vals = [i, p.get("Box No", i), p.get("Part",""), p.get("Brand",""), p.get("Product Details",""), p.get("Length",0), p.get("Breadth",0), p.get("Height",0), p.get("CBM",0), f"{p.get('GW',0)} / {p.get('NW',0)}"]
            for j, v in enumerate(vals):
                row[j].text = str(v)
        summary = docdata.get("packing_summary") or packing_summary(packing)
        document.add_paragraph(
            f"Total Boxes: {int(summary.get('Total Boxes', len(packing)))} | "
            f"Total CBM: {float(summary.get('Total CBM', 0)):.3f} | "
            f"Total GW: {float(summary.get('Total GW', 0)):.2f} KG | "
            f"Total NW: {float(summary.get('Total NW', 0)):.2f} KG"
        )
        if STAMP_PATH.exists():
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run().add_picture(str(STAMP_PATH), width=Inches(1.0))

    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def build_pdf(docdata):
    buffer = BytesIO()
    pdf = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=12*mm, rightMargin=12*mm, topMargin=32*mm, bottomMargin=15*mm)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="Small", parent=styles["Normal"], fontSize=7.5, leading=9))
    styles.add(ParagraphStyle(name="Tiny", parent=styles["Normal"], fontSize=6.8, leading=8))
    styles.add(ParagraphStyle(name="TitleGold", parent=styles["Title"], fontSize=18, textColor=colors.HexColor("#9b763c"), leading=21))
    styles.add(ParagraphStyle(name="Navy", parent=styles["Heading2"], fontSize=11, textColor=colors.HexColor("#071c2e"), leading=13))
    story = []

    story.append(Paragraph(docdata["type"].upper(), styles["TitleGold"]))
    story.append(Paragraph(f"No: {docdata['number']} &nbsp;&nbsp; Date: {docdata['date']} &nbsp;&nbsp; Currency: {docdata['currency']}", styles["Small"]))
    story.append(Paragraph(f"{COMPANY['address']}<br/>{COMPANY['license']}<br/>{COMPANY['phone']}<br/>{COMPANY['email']}", styles["Small"]))
    story.append(Spacer(1,5))

    bill = docdata.get("bill_to",{})
    ship = docdata.get("ship_to",{})
    bill_text = "<br/>".join([f"<b>{bill.get('Company Name','')}</b>", bill.get("Registration Number",""), bill.get("GST/VAT",""), bill.get("Contact Person",""), bill.get("Phone",""), bill.get("Email",""), bill.get("Address",""), bill.get("Country","")])
    ship_text = "Same as Bill To" if docdata.get("ship_same") else "<br/>".join([ship.get("Company Name",""), ship.get("Contact Person",""), ship.get("Phone",""), ship.get("Email",""), ship.get("Address",""), ship.get("Country","")])
    bt = Table([[Paragraph("<b>BILL TO</b><br/>"+bill_text, styles["Small"]), Paragraph("<b>SHIP TO</b><br/>"+ship_text, styles["Small"])]], colWidths=[86*mm,86*mm])
    bt.setStyle(TableStyle([("BOX",(0,0),(-1,-1),0.5,colors.HexColor("#c9b083")),("BACKGROUND",(0,0),(-1,-1),colors.HexColor("#faf7f1")),("VALIGN",(0,0),(-1,-1),"TOP")]))
    story.append(bt); story.append(Spacer(1,6))

    rows = [["SL","Brand","Product Details","Size","Finish","Qty","Rate/PC","Amount"]]
    for i,p in enumerate(docdata["products"], 1):
        qty=float(p.get("Qty",0) or 0); rate=float(p.get("Rate Per Piece",0) or 0)
        rows.append([i, p.get("Brand",""), Paragraph(p.get("Product Details",""), styles["Tiny"]), p.get("Size",""), Paragraph(p.get("Finish",""), styles["Tiny"]), qty, money(rate, docdata["currency"]), money(qty*rate, docdata["currency"])])
    t = Table(rows, repeatRows=1, colWidths=[8*mm,24*mm,43*mm,25*mm,34*mm,10*mm,23*mm,24*mm])
    t.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor("#071c2e")),("TEXTCOLOR",(0,0),(-1,0),colors.white),("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),("GRID",(0,0),(-1,-1),0.25,colors.lightgrey),("FONTSIZE",(0,0),(-1,-1),6.8),("VALIGN",(0,0),(-1,-1),"TOP")]))
    story.append(t); story.append(Spacer(1,6))

    subtotal, disc, shipc, total = calculate(docdata["products"], docdata["discount_type"], docdata["discount_value"], docdata["shipping_enabled"], docdata["shipping_cost"])
    totals = [["Subtotal", money(subtotal, docdata["currency"])],["Discount", f"- {money(disc, docdata['currency'])}"],["Shipping", money(shipc, docdata["currency"])],["Grand Total", money(total, docdata["currency"])]]
    tt = Table(totals, colWidths=[40*mm,38*mm], hAlign="RIGHT")
    tt.setStyle(TableStyle([("GRID",(0,0),(-1,-1),0.25,colors.grey),("BACKGROUND",(0,3),(-1,3),colors.HexColor("#071c2e")),("TEXTCOLOR",(0,3),(-1,3),colors.white),("FONTNAME",(0,3),(-1,3),"Helvetica-Bold"),("FONTSIZE",(0,0),(-1,-1),8)]))
    story.append(tt); story.append(Spacer(1,5))
    story.append(Paragraph(f"<b>Amount in Words:</b> {amount_in_words(total, docdata['currency'])}", styles["Navy"]))
    story.append(Spacer(1,7))

    story.append(Paragraph("BANK DETAILS", styles["Navy"]))
    bank_rows = [["Bank Name","Currency","IBAN","SWIFT/BIC","Bank Address"]]
    for b in BANKS[docdata["currency"]]:
        bank_rows.append([b["Bank Name"], b["Currency"], b["IBAN"], b["SWIFT/BIC"], Paragraph(b["Bank Address"], styles["Tiny"])])
    bank_table = Table(bank_rows, repeatRows=1, colWidths=[38*mm,16*mm,45*mm,25*mm,55*mm])
    bank_table.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor("#9b763c")),("TEXTCOLOR",(0,0),(-1,0),colors.white),("GRID",(0,0),(-1,-1),0.25,colors.grey),("FONTSIZE",(0,0),(-1,-1),6.8),("VALIGN",(0,0),(-1,-1),"TOP")]))
    story.append(bank_table)

    story.append(PageBreak())
    story.append(Paragraph("TERMS & CONDITIONS", styles["TitleGold"]))
    if LOGO_PATH.exists(): story.append(Spacer(1,2))
    for line in docdata.get("terms","").split("\n"):
        if line.strip():
            story.append(Paragraph(line.strip(), styles["Small"]))
            story.append(Spacer(1,1.5))
    story.append(Spacer(1,12))
    sig = [["Seller Signature", "Buyer Signature"]]
    if STAMP_PATH.exists():
        sig.append([Image(str(STAMP_PATH), width=28*mm, height=20*mm), ""])
    sig.append(["HARSH TEJPAL RANA\nOWNER\nCASA ARTE PRIVEE FZ-LLC", "Name:\nCompany:\nDate:"])
    sigt = Table(sig, colWidths=[85*mm,85*mm])
    sigt.setStyle(TableStyle([("GRID",(0,0),(-1,-1),0.25,colors.grey),("BACKGROUND",(0,0),(-1,0),colors.HexColor("#faf7f1")),("VALIGN",(0,0),(-1,-1),"TOP"),("FONTSIZE",(0,0),(-1,-1),8)]))
    story.append(sigt)

    if docdata["type"] == "Invoice":
        story.append(PageBreak())
        story.append(Paragraph("PACKING LIST", styles["TitleGold"]))
        pack = packing_from_products(docdata["products"], docdata.get("packing",[]))
        prow = [["SL","Box No","Part","Brand","Product Details","Length","Breadth","Height","CBM","GW","NW"]]
        for i,p in enumerate(pack, 1):
            prow.append([i,p.get("Box No", i),p.get("Part","1/1"),p["Brand"],Paragraph(p["Product Details"],styles["Tiny"]),p["Length"],p["Breadth"],p["Height"],p["CBM"],p["GW"],p["NW"]])
        pt = Table(prow, repeatRows=1, colWidths=[7*mm,13*mm,13*mm,24*mm,45*mm,13*mm,13*mm,13*mm,16*mm,16*mm,16*mm])
        pt.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor("#071c2e")),("TEXTCOLOR",(0,0),(-1,0),colors.white),("GRID",(0,0),(-1,-1),0.25,colors.lightgrey),("FONTSIZE",(0,0),(-1,-1),6.8),("VALIGN",(0,0),(-1,-1),"TOP")]))
        story.append(pt); story.append(Spacer(1,8))
        summary = docdata.get("packing_summary") or packing_summary(pack)
        story.append(Paragraph(f"<b>Total Boxes:</b> {int(summary.get('Total Boxes', len(pack)))} &nbsp;&nbsp; <b>Total CBM:</b> {float(summary.get('Total CBM', 0)):.3f} &nbsp;&nbsp; <b>Total GW:</b> {float(summary.get('Total GW', 0)):.2f} KG &nbsp;&nbsp; <b>Total NW:</b> {float(summary.get('Total NW', 0)):.2f} KG", styles["Navy"]))
        story.append(Spacer(1,15))
        story.append(Paragraph("Digital Signature / Authorized Signatory", styles["Navy"]))
        if STAMP_PATH.exists():
            story.append(Image(str(STAMP_PATH), width=30*mm, height=22*mm))
        story.append(Paragraph("HARSH TEJPAL RANA<br/>OWNER<br/>CASA ARTE PRIVEE FZ-LLC", styles["Small"]))

    pdf.build(story, onFirstPage=lambda c,d: pdf_header_footer(c,d,docdata["type"]), onLaterPages=lambda c,d: pdf_header_footer(c,d,docdata["type"]))
    return buffer.getvalue()

settings = load_json(SETTINGS_FILE, {"terms":"","password":"1985"})
st.set_page_config(page_title="Casa Arte Privée Billing", layout="wide")

# optional password, on by default
if "auth" not in st.session_state: st.session_state.auth = False
if not st.session_state.auth:
    st.markdown("<div style='background:#071c2e;padding:38px;border-radius:22px;text-align:center;color:white;margin-top:80px;'><h1 style='color:#d0aa65;font-family:Georgia;'>CASA ARTE PRIVÉE</h1><p style='letter-spacing:4px;'>BILLING SYSTEM LOGIN</p></div>", unsafe_allow_html=True)
    pw = st.text_input("Password", type="password")
    if st.button("Login"):
        if pw == str(settings.get("password","1985")):
            st.session_state.auth = True
            st.rerun()
        else:
            st.error("Incorrect password")
    st.stop()

st.markdown("""
<style>
.stApp { background:#f7f3ec; }
.block-container { max-width:1600px; padding-top:1rem; }
div[data-testid="stButton"] button { font-size: 11px !important; padding: 0.25rem 0.35rem !important; min-height: 30px !important; }

.cap { background:#071c2e; color:white; padding:22px; border-radius:22px; margin-bottom:18px; box-shadow:0 12px 30px #0002; }
.gold { color:#d0aa65; font-family:Georgia,serif; }
.card { background:white; border:1px solid #eadfcd; border-radius:18px; padding:16px; margin-bottom:16px; box-shadow:0 8px 24px #071c2e10; }
</style>
""", unsafe_allow_html=True)

logo = img_b64(LOGO_PATH)
logo_html = f"<img src='data:image/png;base64,{logo}' style='height:78px;object-fit:contain;margin-right:18px;'>" if logo else ""
st.markdown(f"<div class='cap' style='display:flex;align-items:center;'>{logo_html}<div><h1 class='gold' style='margin:0;font-size:34px;'>CASA ARTE PRIVÉE</h1><div style='letter-spacing:5px;font-size:11px;'>BILLING · PROFORMA · INVOICE · PACKING LIST</div></div></div>", unsafe_allow_html=True)

customers = load_json(CUSTOMERS_FILE, [])
documents = load_json(DOCUMENTS_FILE, [])

if "editing_id" not in st.session_state: st.session_state.editing_id = None



PAGES = ["Create / Edit", "Saved Documents", "Customers", "Settings"]

if "page" not in st.session_state:
    st.session_state.page = "Create / Edit"

# Hidden navigation intent used by Edit / Convert / Save actions.
if "force_page" in st.session_state:
    st.session_state.page = st.session_state.force_page
    del st.session_state.force_page

# Compact top navigation - no wide sidebar.
nav_cols = st.columns([1, 1, 1, 1, 6])
if nav_cols[0].button("Create / Edit", use_container_width=True):
    st.session_state.page = "Create / Edit"
    st.rerun()
if nav_cols[1].button("Saved", use_container_width=True):
    st.session_state.page = "Saved Documents"
    st.rerun()
if nav_cols[2].button("Customers", use_container_width=True):
    st.session_state.page = "Customers"
    st.rerun()
if nav_cols[3].button("Settings", use_container_width=True):
    st.session_state.page = "Settings"
    st.rerun()

st.caption(f"Current page: {st.session_state.page}")

if st.session_state.page == "Create / Edit":
    editing = next((d for d in documents if d.get("id") == st.session_state.editing_id), None) if st.session_state.editing_id else None
    if editing:
        st.success(f"Editing existing document: {editing.get('number','')}. Saving will update the SAME document.")
    else:
        st.info("Creating a new document.")

    c1,c2,c3,c4 = st.columns(4)
    doc_type = c1.selectbox("Document Type", ["Proforma Invoice","Invoice"], index=1 if editing and editing.get("type")=="Invoice" else 0)
    currency = c2.selectbox("Currency", ["EUR","USD","AED"], index=["EUR","USD","AED"].index(editing.get("currency","EUR")) if editing else 0)
    doc_date = c3.date_input("Date", value=datetime.strptime(editing.get("date"), "%Y-%m-%d").date() if editing else date.today())
    doc_number = c4.text_input("Document No.", value=editing.get("number") if editing else next_number(doc_type, documents))

    if st.button("Start New Blank Document"):
        st.session_state.editing_id = None
        st.session_state.page = "Create / Edit"
        st.rerun()

    st.markdown("<div class='card'><h3 class='gold'>Customer Details</h3>", unsafe_allow_html=True)
    names = ["-- New Customer --"] + [c.get("Company Name","") for c in customers]
    selected = st.selectbox("Choose saved customer", names)
    selected_customer = next((c for c in customers if c.get("Company Name")==selected), {}) if selected != "-- New Customer --" else {}
    bill_existing = editing.get("bill_to", {}) if editing else selected_customer
    a,b = st.columns(2)
    with a:
        bill_company = st.text_input("Company Name", value=bill_existing.get("Company Name",""))
        bill_reg = st.text_input("Company Registration Number", value=bill_existing.get("Registration Number",""))
        bill_vat = st.text_input("GST / VAT", value=bill_existing.get("GST/VAT",""))
        bill_contact = st.text_input("Contact Person", value=bill_existing.get("Contact Person",""))
    with b:
        bill_phone = st.text_input("Phone", value=bill_existing.get("Phone",""))
        bill_email = st.text_input("Email", value=bill_existing.get("Email",""))
        bill_country = st.text_input("Country", value=bill_existing.get("Country",""))
        bill_address = st.text_area("Billing Address", value=bill_existing.get("Address",""), height=92)
    ship_same = st.checkbox("Ship To same as Bill To", value=editing.get("ship_same", True) if editing else True)
    ship_to = {}
    if not ship_same:
        s1,s2 = st.columns(2)
        ship_existing = editing.get("ship_to", {}) if editing else selected_customer.get("ship_to", {})
        with s1:
            ship_to["Company Name"] = st.text_input("Shipping Company", value=ship_existing.get("Company Name",""))
            ship_to["Contact Person"] = st.text_input("Shipping Contact", value=ship_existing.get("Contact Person",""))
            ship_to["Phone"] = st.text_input("Shipping Phone", value=ship_existing.get("Phone",""))
        with s2:
            ship_to["Email"] = st.text_input("Shipping Email", value=ship_existing.get("Email",""))
            ship_to["Country"] = st.text_input("Shipping Country", value=ship_existing.get("Country",""))
            ship_to["Address"] = st.text_area("Shipping Address", value=ship_existing.get("Address",""), height=92)
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("<div class='card'><h3 class='gold'>Products</h3>", unsafe_allow_html=True)
    init_products = editing.get("products") if editing else [{"Brand":"","Product Details":"","Size":"","Finish":"","Qty":1,"Rate Per Piece":0.0}]
    pdf = pd.DataFrame(init_products)
    for col in PRODUCT_COLS:
        if col not in pdf.columns: pdf[col] = 0 if col in ["Qty","Rate Per Piece"] else ""
    edited_df = st.data_editor(pdf[PRODUCT_COLS], num_rows="dynamic", use_container_width=True, key=f"prod_{st.session_state.editing_id or 'new'}")
    products = edited_df.fillna("").to_dict("records")
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("<div class='card'><h3 class='gold'>Discount / Shipping / Totals</h3>", unsafe_allow_html=True)
    d1,d2,d3,d4 = st.columns(4)
    discount_type = d1.selectbox("Discount Type", ["Percentage","Flat Amount"], index=["Percentage","Flat Amount"].index(editing.get("discount_type","Percentage")) if editing else 0)
    discount_value = d2.number_input("Discount Value", min_value=0.0, value=float(editing.get("discount_value",0) if editing else 0))
    shipping_enabled = d3.checkbox("Add Shipping Charges?", value=bool(editing.get("shipping_enabled", False)) if editing else False)
    shipping_cost = d4.number_input("Shipping Cost", min_value=0.0, value=float(editing.get("shipping_cost",0) if editing else 0), disabled=not shipping_enabled, help="Tick Add Shipping Charges first, then enter shipping amount.")
    subtotal, discount_amount, shipping_amount, grand = calculate(products, discount_type, discount_value, shipping_enabled, shipping_cost)
    m1,m2,m3,m4 = st.columns(4)
    m1.metric("Subtotal", money(subtotal, currency))
    m2.metric("Discount", money(discount_amount, currency))
    m3.metric("Shipping", money(shipping_amount, currency))
    m4.metric("Grand Total", money(grand, currency))
    st.info(f"Amount in Words: {amount_in_words(grand, currency)}")
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("<div class='card'><h3 class='gold'>Bank Details</h3>", unsafe_allow_html=True)
    st.dataframe(pd.DataFrame(BANKS[currency]), use_container_width=True, hide_index=True)
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("<div class='card'><h3 class='gold'>Terms & Conditions</h3>", unsafe_allow_html=True)
    terms = st.text_area("Editable default terms for this document", value=editing.get("terms", settings.get("terms","")) if editing else settings.get("terms",""), height=330)
    st.markdown("</div>", unsafe_allow_html=True)

    packing = []
    packing_summary_values = {"Total Boxes": 0, "Total CBM": 0.0, "Total GW": 0.0, "Total NW": 0.0}
    packing_manual_override = False

    if doc_type == "Invoice":
        st.markdown("<div class='card'><h3 class='gold'>Packing List — Mandatory</h3>", unsafe_allow_html=True)

        current_doc_key = editing.get("id") if editing else "new_document"
        packing_state_key = f"packing_rows_{current_doc_key}"

        if st.session_state.get("active_packing_doc_key") != current_doc_key:
            saved_packing = editing.get("packing", []) if editing else []
            st.session_state[packing_state_key] = clean_packing_rows(packing_from_products(products, saved_packing))
            st.session_state["active_packing_doc_key"] = current_doc_key

        if packing_state_key not in st.session_state:
            saved_packing = editing.get("packing", []) if editing else []
            st.session_state[packing_state_key] = clean_packing_rows(packing_from_products(products, saved_packing))

        existing_rows = clean_packing_rows(st.session_state[packing_state_key])
        existing_pairs = {(r.get("Brand", ""), r.get("Product Details", "")) for r in existing_rows}

        # Add one default packing row only for each actual invoice product.
        for p in products:
            if not str(p.get("Brand", "")).strip() and not str(p.get("Product Details", "")).strip():
                continue
            pair = (p.get("Brand", ""), p.get("Product Details", ""))
            if pair not in existing_pairs:
                existing_rows.append({
                    "Box No": len(existing_rows) + 1,
                    "Part": "1/1",
                    "Brand": p.get("Brand", ""),
                    "Product Details": p.get("Product Details", ""),
                    "Length": 0.0,
                    "Breadth": 0.0,
                    "Height": 0.0,
                    "CBM": 0.0,
                    "GW": 0.0,
                    "NW": 0.0,
                })
                existing_pairs.add(pair)

        st.session_state[packing_state_key] = clean_packing_rows(existing_rows)

        st.caption("Use Add/Split Box when an item has 2 or more boxes/parts. Blank rows are ignored and not counted.")

        product_labels = [
            f"{i+1}. {p.get('Brand','')} - {p.get('Product Details','')}"
            for i, p in enumerate(products)
            if str(p.get("Brand", "")).strip() or str(p.get("Product Details", "")).strip()
        ]

        split_cols = st.columns([3, 1, 1])
        with split_cols[0]:
            split_choice = st.selectbox("Select item to add another box/part", product_labels if product_labels else ["No products"])
        with split_cols[1]:
            part_label = st.text_input("Part label", value="2/2")
        with split_cols[2]:
            st.write("")
            st.write("")
            if st.button("Add/Split Box"):
                if product_labels:
                    original_index = int(split_choice.split(".")[0]) - 1
                    p = products[original_index]
                    st.session_state[packing_state_key].append({
                        "Box No": len(clean_packing_rows(st.session_state[packing_state_key])) + 1,
                        "Part": part_label or "Part",
                        "Brand": p.get("Brand", ""),
                        "Product Details": p.get("Product Details", ""),
                        "Length": 0.0,
                        "Breadth": 0.0,
                        "Height": 0.0,
                        "CBM": 0.0,
                        "GW": 0.0,
                        "NW": 0.0,
                    })
                    st.session_state[packing_state_key] = clean_packing_rows(st.session_state[packing_state_key])
                    st.rerun()

        pack_init = pd.DataFrame(clean_packing_rows(st.session_state[packing_state_key]))
        for col in PACK_COLS:
            if col not in pack_init.columns:
                pack_init[col] = 0 if col in ["Box No","Length","Breadth","Height","CBM","GW","NW"] else ""

        pack_edit = st.data_editor(
            pack_init[PACK_COLS],
            num_rows="dynamic",
            use_container_width=True,
            key=f"pack_editor_{current_doc_key}"
        )

        packing = pack_edit.fillna("").to_dict("records")
        real_packing = []
        for row in packing:
            if not is_real_packing_row(row):
                continue
            row["CBM"] = round(
                float(row.get("Length", 0) or 0)
                * float(row.get("Breadth", 0) or 0)
                * float(row.get("Height", 0) or 0)
                / 1000000,
                3
            )
            row["GW"] = float(row.get("GW", 0) or 0)
            row["NW"] = float(row.get("NW", 0) or 0)
            real_packing.append(row)

        packing = clean_packing_rows(real_packing)
        st.session_state[packing_state_key] = packing

        auto_summary = packing_summary(packing)

        st.markdown("**Packing Summary**")
        packing_manual_override = st.checkbox("Manual override summary totals", value=False)
        sc1, sc2, sc3, sc4 = st.columns(4)

        if packing_manual_override:
            total_boxes = sc1.number_input("Total Boxes", min_value=0, value=int(auto_summary["Total Boxes"]))
            total_cbm = sc2.number_input("Total CBM", min_value=0.0, value=float(auto_summary["Total CBM"]), format="%.3f")
            total_gw = sc3.number_input("Total GW", min_value=0.0, value=float(auto_summary["Total GW"]), format="%.2f")
            total_nw = sc4.number_input("Total NW", min_value=0.0, value=float(auto_summary["Total NW"]), format="%.2f")
        else:
            total_boxes = auto_summary["Total Boxes"]
            total_cbm = auto_summary["Total CBM"]
            total_gw = auto_summary["Total GW"]
            total_nw = auto_summary["Total NW"]
            sc1.metric("Total Boxes", total_boxes)
            sc2.metric("Total CBM", f"{total_cbm:.3f}")
            sc3.metric("Total GW", f"{total_gw:.2f} KG")
            sc4.metric("Total NW", f"{total_nw:.2f} KG")

        packing_summary_values = {
            "Total Boxes": int(total_boxes),
            "Total CBM": float(total_cbm),
            "Total GW": float(total_gw),
            "Total NW": float(total_nw),
            "Manual Override": bool(packing_manual_override),
        }

        st.markdown("</div>", unsafe_allow_html=True)
    else:
        st.warning("Proforma: packing list is hidden. It will auto-create when converted to Invoice.")

    bill_to = {"Company Name":bill_company,"Registration Number":bill_reg,"GST/VAT":bill_vat,"Contact Person":bill_contact,"Phone":bill_phone,"Email":bill_email,"Country":bill_country,"Address":bill_address}
    docdata = {
        "id": editing.get("id") if editing else str(uuid.uuid4()),
        "type": doc_type, "number": doc_number, "date": str(doc_date), "currency": currency,
        "bill_to": bill_to, "ship_same": ship_same, "ship_to": ship_to,
        "products": products, "discount_type": discount_type, "discount_value": discount_value,
        "shipping_enabled": shipping_enabled, "shipping_cost": shipping_cost,
        "terms": terms, "packing": packing, "packing_summary": packing_summary_values, "total": grand,
        "created_at": editing.get("created_at") if editing else datetime.now().isoformat(timespec="seconds"),
        "updated_at": datetime.now().isoformat(timespec="seconds")
    }

    x1,x2,x3 = st.columns(3)
    if x1.button("Save / Update Document", type="primary"):
        # save/update customer
        if bill_company:
            cust = dict(bill_to)
            cust["ship_to"] = bill_to if ship_same else ship_to
            found = next((i for i,c in enumerate(customers) if c.get("Company Name","").lower()==bill_company.lower()), None)
            if found is None:
                customers.append(cust)
            else:
                customers[found] = cust
            save_json(CUSTOMERS_FILE, customers)

        # IMPORTANT: if editing existing document, update the same record by ID.
        # If creating new, append only once.
        existing_id = editing.get("id") if editing else docdata.get("id")
        docdata["id"] = existing_id

        idx = next((i for i,d in enumerate(documents) if d.get("id") == existing_id), None)
        if idx is None:
            documents.append(docdata)
        else:
            documents[idx] = docdata

        save_json(DOCUMENTS_FILE, documents)

        # After save/update, return to Saved Documents list.
        st.session_state.editing_id = None
        st.session_state.force_page = "Saved Documents"
        st.success(f"Saved / Updated: {docdata['number']}")
        st.rerun()

    if x2.button("Convert Proforma to Invoice"):
        if doc_type == "Proforma Invoice":
            newdoc = dict(docdata)
            newdoc["id"] = str(uuid.uuid4())
            newdoc["type"] = "Invoice"
            newdoc["number"] = next_number("Invoice", documents)
            newdoc["packing"] = packing_from_products(products)
            newdoc["created_at"] = datetime.now().isoformat(timespec="seconds")
            newdoc["updated_at"] = datetime.now().isoformat(timespec="seconds")
            documents.append(newdoc)
            save_json(DOCUMENTS_FILE, documents)
            st.session_state.editing_id = newdoc["id"]
            st.session_state.page = "Create / Edit"
            st.success(f"Converted to Invoice: {newdoc['number']}. Opening converted invoice for editing/packing details...")
            st.rerun()
        else:
            st.info("Already an Invoice.")

    x3.download_button("Download PDF", data=build_pdf(docdata), file_name=f"{doc_number.replace('/','-')}.pdf", mime="application/pdf")
    st.download_button("Download Word", data=build_word(docdata), file_name=f"{doc_number.replace('/','-')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if st.session_state.page == "Saved Documents":
    st.markdown("<div class='card'><h3 class='gold'>Saved Documents — Select / Edit / Convert / Delete</h3>", unsafe_allow_html=True)
    search = st.text_input("Search saved documents by number, customer, type, currency")
    results = documents
    if search:
        s = search.lower()
        results = [d for d in documents if s in json.dumps(d, ensure_ascii=False).lower()]

    if not results:
        st.warning("No documents saved.")
    else:
        st.caption("Each document has its own Edit, Convert, Download and Delete action. Edit/Convert will take you back to the data entry page.")
        for d in sorted(results, key=lambda x: x.get("updated_at", x.get("created_at","")), reverse=True):
            customer_name = d.get("bill_to", {}).get("Company Name", "")
            row = st.container(border=True)
            with row:
                c0, c1, c2, c3, c4, c5 = st.columns([2.4, 1.3, 2.2, 1.1, 1.2, 1.8])
                c0.markdown(f"### {d.get('number','')}")
                c0.caption(f"{d.get('date','')} · {d.get('currency','')}")
                c1.markdown(f"**{d.get('type','')}**")
                c2.markdown(customer_name or "No customer")
                c3.markdown(f"**{money(d.get('total',0), d.get('currency','EUR'))}**")

                if c4.button("Edit", key=f"edit_{d.get('id')}"):
                    st.session_state.editing_id = d.get("id")
                    st.session_state.force_page = "Create / Edit"
                    st.rerun()

                if d.get("type") == "Proforma Invoice":
                    if c5.button("Convert to Invoice", key=f"convert_{d.get('id')}"):
                        newdoc = dict(d)
                        newdoc["id"] = str(uuid.uuid4())
                        newdoc["type"] = "Invoice"
                        newdoc["number"] = next_number("Invoice", documents)
                        newdoc["packing"] = packing_from_products(newdoc.get("products", []))
                        newdoc["created_at"] = datetime.now().isoformat(timespec="seconds")
                        newdoc["updated_at"] = datetime.now().isoformat(timespec="seconds")
                        documents.append(newdoc)
                        save_json(DOCUMENTS_FILE, documents)
                        st.session_state.editing_id = newdoc["id"]
                        st.session_state.force_page = "Create / Edit"
                        st.rerun()
                else:
                    c5.caption("Already invoice")

                d1, d2 = st.columns([1.2, 2])
                d1.download_button(
                    "Download PDF",
                    data=build_pdf(d),
                    file_name=f"{d.get('number','document').replace('/','-')}.pdf",
                    mime="application/pdf",
                    key=f"pdf_{d.get('id')}"
                )
                d1.download_button(
                    "Download Word",
                    data=build_word(d),
                    file_name=f"{d.get('number','document').replace('/','-')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"docx_{d.get('id')}"
                )

                confirm = d2.checkbox(f"Confirm delete {d.get('number','')}", key=f"confirm_delete_{d.get('id')}")
                if d2.button("Delete", key=f"delete_{d.get('id')}", type="secondary"):
                    if confirm:
                        documents = [x for x in documents if x.get("id") != d.get("id")]
                        save_json(DOCUMENTS_FILE, documents)
                        if st.session_state.editing_id == d.get("id"):
                            st.session_state.editing_id = None
                        st.session_state.force_page = "Saved Documents"
                        st.success(f"Deleted {d.get('number','document')}")
                        st.rerun()
                    else:
                        st.warning("Tick confirm delete first.")
    st.markdown("</div>", unsafe_allow_html=True)

if st.session_state.page == "Customers":
    st.markdown("<div class='card'><h3 class='gold'>Customer Database</h3>", unsafe_allow_html=True)
    if customers:
        st.dataframe(pd.DataFrame(customers).drop(columns=["ship_to"], errors="ignore"), use_container_width=True, hide_index=True)
        customer_names_for_delete = [c.get("Company Name","") for c in customers if c.get("Company Name","")]
        if customer_names_for_delete:
            selected_customer_delete = st.selectbox("Select customer to delete", customer_names_for_delete)
            confirm_delete_customer = st.checkbox(f"Confirm delete customer {selected_customer_delete}")
            if st.button("Delete Selected Customer", type="secondary"):
                if confirm_delete_customer:
                    customers = [c for c in customers if c.get("Company Name","") != selected_customer_delete]
                    save_json(CUSTOMERS_FILE, customers)
                    st.success(f"Deleted customer: {selected_customer_delete}")
                    st.rerun()
                else:
                    st.warning("Please tick confirm delete first.")
    else:
        st.info("No customers saved yet.")
    st.markdown("</div>", unsafe_allow_html=True)

if st.session_state.page == "Settings":
    st.markdown("<div class='card'><h3 class='gold'>Settings</h3>", unsafe_allow_html=True)
    st.write("Password is currently set to **1985**.")
    new_terms = st.text_area("Default Terms & Conditions", value=settings.get("terms",""), height=430)
    if st.button("Save Default Terms"):
        settings["terms"] = new_terms
        save_json(SETTINGS_FILE, settings)
        st.success("Settings saved.")
    st.markdown("</div>", unsafe_allow_html=True)
